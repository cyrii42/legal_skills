#!/usr/bin/env python3
# /// script
# requires-python = ">=3.9"
# dependencies = ["python-docx>=1.1"]
# ///
"""Build a house-styled .docx deposition summary from a Markdown file.

House style: Times New Roman 12pt; bold, section-numbered headings at body
size; 1-inch margins; footer page numbers; left-aligned, single-spaced body;
a "Privileged & Confidential - Attorney Work Product" legend at the top.

Handles the constrained Markdown subset the deposition-summary skill emits:
optional YAML frontmatter, ATX headings (#/##/###), paragraphs, unordered and
ordered lists (nested by indentation), pipe tables, and block quotes. Inline
**bold**, *italic*, and `code` are rendered.

Usage:
    uv run --script build_docx.py <summary.md> [output.docx]
    # or, in an environment with python-docx already available:
    python3 build_docx.py <summary.md> [output.docx]
    # or import and call build(md_path, out_path)
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT = "Times New Roman"
SIZE = Pt(12)
LEGEND = "Privileged & Confidential - Attorney Work Product"

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")

# Curly-quote characters.
_LDQUO, _RDQUO, _LSQUO, _RSQUO = "“", "”", "‘", "’"
# A straight quote is "opening" when it starts the text or follows whitespace or
# one of these opening delimiters (including an em/en-dash or an opening curly
# double quote, so a nested single quote is treated as opening).
_OPEN_AFTER = set(" \t\n\r([{<—–")
_CODE_SPAN = re.compile(r"(```.*?```|`[^`\n]*`)", re.DOTALL)
_WIKILINK = re.compile(r"!?\[\[([^\]|]+)(?:\|([^\]]+))?\]\]")


def _dewiki(text):
    """Render Obsidian wikilinks/embeds as plain readable text. A Word document
    has no concept of them, so [[a/b#c|Display]] -> Display and a bare
    [[a/b#c]] -> "b" (path and #anchor stripped). Leaves code spans alone."""
    def repl(m):
        if m.group(2):
            return m.group(2)
        return m.group(1).split("#", 1)[0].rsplit("/", 1)[-1]

    parts = _CODE_SPAN.split(text)
    for i in range(0, len(parts), 2):
        parts[i] = _WIKILINK.sub(repl, parts[i])
    return "".join(parts)


def _educate_plain(s):
    """Turn straight quotes into typographic quotes using SmartyPants-style
    heuristics. Apostrophes and closing quotes default to the right-hand mark."""
    out = []
    for i, ch in enumerate(s):
        prev = s[i - 1] if i else "\n"
        if ch == '"':
            out.append(_LDQUO if prev in _OPEN_AFTER else _RDQUO)
        elif ch == "'":
            out.append(_LSQUO if prev in _OPEN_AFTER or prev in '"' + _LDQUO
                       else _RSQUO)
        else:
            out.append(ch)
    return "".join(out)


def _educate(text):
    """Educate quotes across the document, leaving inline/fenced code spans
    (Bates labels, Westlaw asterisk cites, etc.) untouched."""
    parts = _CODE_SPAN.split(text)
    for i in range(0, len(parts), 2):  # even indices are outside code spans
        parts[i] = _educate_plain(parts[i])
    return "".join(parts)


def _style_run(run):
    run.font.name = FONT
    run.font.size = SIZE
    # Ensure the East Asian / complex-script slots also use the face, so Word
    # does not silently substitute a different font.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def _add_runs(paragraph, text):
    """Add inline-formatted runs to a paragraph."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _style_run(paragraph.add_run(text[pos:m.start()]))
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
        else:  # `code`
            r = paragraph.add_run(tok[1:-1])
        _style_run(r)
        pos = m.end()
    if pos < len(text):
        _style_run(paragraph.add_run(text[pos:]))


def _add_page_number(paragraph):
    run = paragraph.add_run()
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)
    _style_run(run)


def _setup(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SIZE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # Confidentiality legend as a right-justified page header, repeating on
        # every page (not a body paragraph).
        section.different_first_page_header_footer = False
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.text = ""
        hr = header_p.add_run(LEGEND)
        hr.bold = True
        _style_run(hr)
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(footer_p)


def _heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for run in p.runs:
        run.bold = True
        if level >= 3:
            run.italic = True
    return p


def _list_style(ordered, level):
    base = "List Number" if ordered else "List Bullet"
    return base if level == 0 else f"{base} {min(level + 1, 3)}"


def _leading_indent(raw):
    """Return nesting level from leading whitespace (tab or 2+ spaces = 1)."""
    m = re.match(r"[ \t]*", raw)
    ws = m.group(0)
    tabs = ws.count("\t")
    spaces = len(ws) - tabs
    return tabs + spaces // 2


def _table(doc, rows):
    # rows: list of lists of cell strings; row 0 is the header.
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _add_runs(p, row[j] if j < len(row) else "")
            if i == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def _blockquote(doc, lines):
    for line in lines:
        text = line.lstrip(">").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(2)
        if text:
            _add_runs(p, text)


def _split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", line)) and "-" in line


def build(md_path, out_path=None):
    md_path = Path(md_path)
    if out_path is None:
        out_path = md_path.with_suffix(".docx")
    out_path = Path(out_path)

    text = md_path.read_text(encoding="utf-8")
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = _dewiki(text)  # Obsidian wikilinks -> plain text (meaningless in Word)
    text = _educate(text)  # straight quotes -> typographic quotes (house style)
    lines = text.splitlines()

    # Strip YAML frontmatter.
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                lines = lines[i + 1:]
                break

    doc = Document()
    _setup(doc)

    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # Skip a markdown legend line if the model included one (avoid dupe).
        if "privileged & confidential" in line.lower():
            i += 1
            continue

        # Headings.
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            _heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # Horizontal rule: skip.
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", line):
            i += 1
            continue

        # Tables (header line followed by a separator line).
        if line.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            rows = [_split_row(line)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            _table(doc, rows)
            continue

        # Block quotes.
        if line.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i])
                i += 1
            _blockquote(doc, block)
            continue

        # Lists.
        ul = re.match(r"^[-*]\s+(.*)$", line)
        ol = re.match(r"^\d+[.)]\s+(.*)$", line)
        if ul or ol:
            ordered = bool(ol)
            content = (ol or ul).group(1)
            level = _leading_indent(raw)
            p = doc.add_paragraph(style=_list_style(ordered, level))
            p.paragraph_format.space_after = Pt(2)
            _add_runs(p, content)
            i += 1
            continue

        # Paragraph: gather following non-blank, non-structural lines.
        para = [line]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or re.match(r"^(#{1,6})\s", nxt) or nxt.startswith(">")
                    or nxt.startswith("|") or re.match(r"^[-*]\s+", nxt)
                    or re.match(r"^\d+[.)]\s+", nxt)
                    or re.match(r"^(-{3,}|\*{3,}|_{3,})$", nxt)):
                break
            para.append(nxt)
            i += 1
        p = doc.add_paragraph()
        _add_runs(p, " ".join(para))

    doc.save(str(out_path))
    return str(out_path)


def main(argv):
    if len(argv) < 2:
        print("usage: build_docx.py <summary.md> [output.docx]", file=sys.stderr)
        return 2
    out = build(argv[1], argv[2] if len(argv) > 2 else None)
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
